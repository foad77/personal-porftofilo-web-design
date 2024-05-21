from docx import Document

# Create a new Document
doc = Document()

# Title
doc.add_heading('Delivery Receipt', level=1)

# Date
doc.add_paragraph('**Date:** 05/20/2024')

# To Section
doc.add_paragraph('**To:**\nAdone Renewables\n[Adone Renewables Address]\n[City, State, Zip Code]')

# Subject
doc.add_paragraph('**Subject:** Delivery of Mobile Phones')

# Items Delivered
doc.add_heading('Items Delivered:', level=2)

# iPhone Details
doc.add_paragraph('1. **iPhone**\n   - **Model:** iPhone SE\n   - **Serial Number:** GWTHGLXUPLMJ\n   - **IMEI:** 355143720257798\n   - **IMEI2:** 355143720678373')

# Android Phone Details
doc.add_paragraph('2. **Android Phone**\n   - **Model:** Galaxy A20\n   - **Model Number:** SM-A20SUI\n   - **Serial Number:** R58MA6S1W9B\n   - **IMEI:** 357190102329792')

# Confirmation
doc.add_paragraph('**Confirmation:**\nI, the undersigned, acknowledge receipt of the above-mentioned items in good condition. I confirm that the serial numbers, IMEI, and model numbers match the phone specifications provided above.')

# Received By Section
doc.add_paragraph('**Received By:**\n\n**Name:** _______________________________\n\n**Date:** _______________________________\n\n**Signature:** ___________________________')

# Delivery Personnel Details Section
doc.add_paragraph('**Delivery Personnel Details:**\n\n**Name:** _______________________________\n\n**Date:** _______________________________\n\n**Signature:** ___________________________')

# Save the document
file_path = '/mnt/data/Delivery_Receipt.docx'
doc.save(file_path)

file_path
